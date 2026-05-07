"""Word 文档构建器模块"""

import logging
import os
import re
import subprocess
import tempfile
from copy import deepcopy
from dataclasses import dataclass
from typing import Optional

import requests
from docx import Document
from docx.text.paragraph import Paragraph
from lxml.etree import _Element as Element
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK

from ..preprocessor import TextBlock, InlineContent
from ..config import DocumentStyleConfig, get_config
from ..exceptions import ExportError
from .latex_converter import LaTeXConverter


@dataclass
class DocumentConfig:
    """Word 文档全局样式配置"""
    title: str = "豆包聊天记录"  # 文档标题
    author: str = "Doubao Export"  # 文档作者
    margin_left: float = 1.0  # 左边距（英寸）
    margin_right: float = 1.0  # 右边距（英寸）
    margin_top: float = 1.0  # 上边距（英寸）
    margin_bottom: float = 1.0  # 下边距（英寸）
    font_name: str = "微软雅黑"  # 正文字体
    font_size: int = 12  # 字号
    style_config: DocumentStyleConfig | None = None  # 样式配置

    def __post_init__(self):
        if self.style_config is None:
            self.style_config = get_config().document_style


class DocxBuilder:
    """
    Word 文档构建器

    将 TextBlock 列表转换为 Word 文档。
    """

    def __init__(self, config: Optional[DocumentConfig] = None) -> None:
        """
        初始化文档构建器

        Args:
            config: 文档配置，为 None 时使用默认配置
        """
        self.config = config or DocumentConfig(style_config=get_config().document_style)
        self.document = Document()  # python-docx 文档对象
        self.latex_converter = LaTeXConverter()  # LaTeX 公式转换器
        self._setup_document()

        # 列表状态追踪（用于有序列表序号连续性）
        self._last_list_type: str | None = None  # 上一个列表类型（"ul" 或 "ol"）
        self._last_list_level: int = 0  # 上一个列表的嵌套层级
        self._list_counter: int = 0  # 有序列表当前序号

        # 图片下载失败追踪
        self._image_failure_count: int = 0  # 失败图片数量
        self._image_failure_urls: list[str] = []  # 失败图片 URL 列表
        self._image_failure_urls_set: set[str] = set()  # 用于去重

        self._config = get_config()
        self._pandoc_timeout: int = self._config.pandoc.timeout if self._config.pandoc else 15
        crawler_config = self._config.crawler
        self._image_timeout: int = crawler_config.image_download_timeout if crawler_config else 15
        self._logger = logging.getLogger(__name__)  # 模块级 logger

    def _setup_document(self) -> None:
        """设置页边距和默认字体"""
        section = self.document.sections[0]
        section.top_margin = Inches(self.config.margin_top)
        section.bottom_margin = Inches(self.config.margin_bottom)
        section.left_margin = Inches(self.config.margin_left)
        section.right_margin = Inches(self.config.margin_right)

        style = self.document.styles["Normal"]
        font = getattr(style, 'font', None)
        if font:
            font.name = self.config.font_name
            font.size = Pt(self.config.font_size)

    def _set_run_font(self, run) -> None:
        """
        设置 Run 字体

        Args:
            run: docx Run 对象
        """
        run.font.name = self.config.font_name
        run.font.size = Pt(self.config.font_size)
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), self.config.font_name)
        except (AttributeError, TypeError):
            pass

    def _add_title(self, title: str) -> None:
        """
        添加文档标题

        Args:
            title: 标题文本
        """
        heading = self.document.add_heading(title, level=1)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in heading.runs:
            run.font.size = Pt(self.config.style_config.title_font_size if self.config.style_config else 18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)
            self._set_run_font(run)

    def _add_paragraph(self, text: str, style: Optional[str] = None) -> None:
        """
        添加普通段落

        Args:
            text: 段落文本
            style: 段落样式
        """
        para = self.document.add_paragraph(text, style=style)
        for run in para.runs:
            self._set_run_font(run)

    def _add_error_paragraph(self, text: str = "[内容解析失败]") -> None:
        """
        添加错误提示段落（红色字体）

        Args:
            text: 错误提示文本
        """
        para = self.document.add_paragraph()
        run = para.add_run(text)
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)  # 红色
        self._set_run_font(run)

    def build_blocks(
        self,
        title: str,
        blocks: list[tuple[str, "TextBlock"]],
        output_path: str,
    ) -> str:
        """
        从内容块构建文档

        Args:
            title: 文档标题
            blocks: (角色, TextBlock) 元组列表
            output_path: 输出文件路径

        Returns:
            str: 输出文件路径
        """
        self._add_title(title)

        # 重置列表状态
        self._list_counter = 0
        self._last_list_type = None
        self._last_list_level = 0

        current_role = ""
        for role, block in blocks:
            # 切换角色时添加角色标签
            if role != current_role:
                current_role = role
                # 角色切换时重置列表状态，确保每条消息的列表序号独立
                self._list_counter = 0
                self._last_list_type = None
                self._last_list_level = 0
                role_label = "用户" if role == "user" else "豆包"
                heading = self.document.add_heading(f"{role_label}", level=3)
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    self._set_run_font(run)

            try:
                self._add_text_block(block)
            except Exception:
                self._logger.exception(f"处理内容块失败 (type={block.type})")
                self._add_error_paragraph()

        footer_mark = self._config.document_style.footer_mark
        if footer_mark:
            self._add_paragraph(footer_mark)

        self.document.save(output_path)
        return output_path

    def get_image_failures(self) -> tuple[int, list[str]]:
        """
        获取图片下载失败信息

        Returns:
            tuple[int, list[str]]: (失败数量, 失败 URL 列表)
        """
        return self._image_failure_count, self._image_failure_urls

    def _add_text_block(self, block: "TextBlock") -> None:
        """
        根据块类型分发处理

        Args:
            block: TextBlock 内容块
        """
        # 公式（允许 content 为空，如某些占位公式由 _add_latex 处理）
        if block.type == "latex":
            if block.content.strip():
                self._add_latex(block.content, is_display=(block.language == "display"))
        # 代码块
        elif block.type == "code":
            if block.content.strip():
                self._add_code_block(block.content, block.language)
        # 标题
        elif block.type == "heading":
            if block.content.strip():
                level = int(block.language) if block.language else 3  # 标题级别
                heading = self.document.add_heading(block.content, level=min(level, 6))
                for run in heading.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    self._set_run_font(run)
                # 标题意味着新的逻辑分组，重置列表状态
                self._last_list_type = None
                self._last_list_level = 0
                self._list_counter = 0
        # 简单列表（已废弃，列表均以 list_item 形式处理）
        elif block.type == "list":
            pass
        # 列表项
        elif block.type == "list_item":
            self._add_list_item(block)
        # 表格
        elif block.type == "table":
            self._add_table(block.data)
        # 引用
        elif block.type == "blockquote":
            if block.content.strip():
                self._add_blockquote(block.content)
        # 图片
        elif block.type == "image":
            if block.content.strip():
                self._add_image(block.content)
        # 内联段落（包含公式/加粗等）
        elif block.type == "paragraph" and block.items:
            self._add_inline_content(block.items)
        # 普通段落
        else:
            content = block.content.rstrip("\n") if block.content else ""
            if content:
                self._add_paragraph(content)

    def _add_list_item(self, block: "TextBlock") -> None:
        """
        添加列表项

        Args:
            block: TextBlock 内容块
        """
        list_type = block.language if block.language in ("ul", "ol") else "ul"  # 列表类型
        level = getattr(block, 'level', 0)  # 嵌套层级

        if block.items:
            # 包含内联内容的列表项
            if list_type == "ol":
                # 记录状态，由解析器携带的 list_start 驱动序号重置
                self._last_list_type = "ol"
                self._last_list_level = level
                # 解析器已保证有序列表项携带 list_marker 和 list_start，直接传递即可
            else:
                self._last_list_type = "ul"
                self._last_list_level = level
            self._add_inline_content(block.items, list_type, level)
        else:
            # 纯文本列表项（无内联内容）
            if list_type == "ol":
                # 记录状态，由解析器携带的 list_start 驱动序号重置
                self._last_list_type = "ol"
                self._last_list_level = level

                # 构造带标记的占位符，并传入 list_start
                placeholder = InlineContent(
                    type="text",
                    content=block.content or "",
                    list_marker="ol",
                    level=level
                )
                if block.list_start is not None:
                    placeholder.list_start = block.list_start
                self._add_inline_content([placeholder], list_type, level)
            else:
                self._last_list_type = "ul"
                self._last_list_level = level
                para = self.document.add_paragraph(block.content, style="List Bullet")
                for run in para.runs:
                    self._set_run_font(run)

    def _handle_list_marker_item(self, item: InlineContent, level: int) -> Paragraph:
        """
        处理带列表标记的文本项，返回创建的段落供后续文本项复用

        Args:
            item: 内联文本项
            level: 嵌套层级

        Returns:
            Paragraph: 创建的列表段落
        """
        para = self.document.add_paragraph(style="List Bullet")
        run = para.add_run(item.content)
        self._set_run_font(run)
        if item.bold:
            run.font.bold = True
        if item.italic:
            run.font.italic = True
        item_level = getattr(item, 'level', level)
        if item_level > 0:
            para.paragraph_format.left_indent = Inches(item_level * 0.5)
        return para

    def _handle_inline_text_item(self, para, item: InlineContent, last_run, last_was_break: bool) -> tuple:
        """
        处理内联文本项

        Args:
            para: 当前段落
            item: 文本项
            last_run: 上一个 Run 对象
            last_was_break: 上一个是否为换行符

        Returns:
            tuple: (last_run, last_was_break, need_new_para)
        """
        text = item.content
        need_new_para = False

        # 情况1：内容为纯空白（仅换行/空格等无实质字符）
        if not text.strip():
            # 纯换行符项：始终添加一个软换行
            if last_run is not None:
                last_run.add_break(WD_BREAK.LINE)
            else:
                # 没有可依附的 run，先添加一个空 run 再 break
                last_run = para.add_run("")
                last_run.add_break(WD_BREAK.LINE)
            return last_run, True, False

        # 情况2：有实际文本内容
        last_was_break = False
        parts = text.split("\n")
        for i, part in enumerate(parts):
            if i > 0 and last_run:
                last_run.add_break(WD_BREAK.LINE)
            if part:
                run = para.add_run(part)
                self._set_run_font(run)
                if item.bold:
                    run.font.bold = True
                if item.italic:
                    run.font.italic = True
                last_run = run
        return last_run, last_was_break, False

    def _handle_inline_image_item(self, url: str, level: int = 0) -> tuple:
        """
        处理内联图片项

        Args:
            url: 图片 URL
            level: 嵌套层级（用于缩进）

        Returns:
            tuple: (need_new_para, para, last_run)
        """
        image_data = self._download_image(url)
        if image_data:
            temp_path = None
            try:
                with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
                    f.write(image_data)
                    temp_path = f.name
                max_width = Inches(self.config.style_config.inline_image_width if self.config.style_config else 4.0)
                para = self.document.add_paragraph()
                if level > 0:
                    para.paragraph_format.left_indent = Inches(level * 0.5)
                para.add_run().add_picture(temp_path, width=max_width)
            except Exception:
                self._record_image_failure(url)
            finally:
                if temp_path and os.path.exists(temp_path):
                    os.unlink(temp_path)
        else:
            self._record_image_failure(url or "(空URL)")
            para = self.document.add_paragraph("[图片加载失败]")
            if level > 0:
                para.paragraph_format.left_indent = Inches(level * 0.5)
        return True, para, None

    def _handle_inline_table_item(self, item: InlineContent, level: int = 0) -> tuple:
        """处理内联表格项"""
        self._add_table(item.data)
        return True, None, None

    def _handle_inline_code_item(self, item: InlineContent, level: int = 0) -> tuple:
        """处理内联代码块项"""
        self._add_code_block(item.content, level=level)
        return True, None, None

    def _add_inline_code_run(self, para, item: InlineContent) -> None:
        """在当前段落中添加行内代码 run（等宽字体）"""
        run = para.add_run(item.content)
        run.font.size = Pt(self.config.style_config.code_font_size if self.config.style_config else 10)
        self._set_code_font(run)
        if item.bold:
            run.font.bold = True
        if item.italic:
            run.font.italic = True

    def _add_inline_content(self, items: list[InlineContent], list_type: Optional[str] = None, level: int = 0) -> None:
        """
        添加内联内容（文本和公式混合的段落）

        展示公式、表格、代码块处理后不立即创建新段落，仅标记 need_new_para。
        下一次遇到文本或行内公式时才创建续写段落，避免尾部空白段落。
        """
        if not items:
            return

        # 删除 items 末尾的纯空白项（换行/空格），避免段落尾部多余空行
        while items and items[-1].type == "text" and not items[-1].content.strip():
            items.pop()
        if not items:
            return

        # 过滤掉紧邻"另起新段元素"前的换行符，避免多余空行
        items = self._filter_unnecessary_linebreaks(items)
        if not items:
            return

        para = None  # 当前段落，非列表情况下可能为 None
        new_para = self._create_inline_paragraph(list_type, level)
        if new_para is not None:
            para = new_para

        last_run = None  # 上一个 Run 对象（用于换行处理）
        need_new_para = False  # 标记是否需要在下一个可写内容前创建续写段落
        last_was_break = False  # 上一个添加的内容是否为换行符（用于合并连续换行）
        current_item_level = level  # 初始层级

        for item in items:
            # 更新当前层级（item 有 level 时）
            if item.level > 0:
                current_item_level = item.level

            # 有列表标记时，创建新的列表段落
            if item.list_marker and item.type == "text":
                if item.list_marker == "ol":  # 有序列表：自行生成序号
                    # 兜底：如果 item 仍携带 list_start（未经过 _add_list_item 清除），则重置计数器
                    if item.list_start is not None:
                        self._list_counter = item.list_start - 1
                    self._list_counter += 1
                    seq = self._list_counter
                    para = self.document.add_paragraph()
                    # 判断内容是否为空（去除空白后）
                    if item.content.strip():
                        run = para.add_run(f"{seq}. {item.content}")
                    else:
                        run = para.add_run(f"{seq}.")
                    self._set_run_font(run)
                    if item.bold:
                        run.font.bold = True
                    if item.italic:
                        run.font.italic = True
                    if current_item_level > 0:
                        para.paragraph_format.left_indent = Inches(current_item_level * 0.5)
                else:  # 无序列表：沿用原逻辑
                    para = self._handle_list_marker_item(item, current_item_level)
                last_run = None
                need_new_para = False
                last_was_break = False
                continue

            # 文本内容
            if item.type == "text":
                if need_new_para or para is None:
                    if need_new_para:
                        para = self._create_continuation_paragraph(list_type, level)
                        need_new_para = False
                    else:
                        para = self.document.add_paragraph()
                        if level > 0:
                            para.paragraph_format.left_indent = Inches(level * 0.5)
                    last_run = None
                    last_was_break = False
                last_run, last_was_break, _ = self._handle_inline_text_item(
                    para, item, last_run, last_was_break
                )
                continue

            # 展示公式（块级）
            if item.type == "latex" and item.is_display:
                if list_type is not None:
                    # 在列表上下文里，块级公式也放到带列表样式的续写段落中，避免分离
                    para = self._create_continuation_paragraph(list_type, level)
                    self._add_latex_to_paragraph(para, item.content, is_display=False)
                else:
                    self._add_latex(item.content, is_display=True)
                need_new_para = True
                last_run = None
                continue

            # 行内公式
            if item.type == "latex":
                if need_new_para or para is None:
                    if need_new_para:
                        para = self._create_continuation_paragraph(list_type, level)
                        need_new_para = False
                    else:
                        para = self.document.add_paragraph()
                        if level > 0:
                            para.paragraph_format.left_indent = Inches(level * 0.5)
                    last_run = None
                    last_was_break = False
                self._add_latex_to_paragraph(para, item.content, is_display=False)
                last_run = None
                continue

            # 图片
            if item.type == "image" and item.image_url:
                need_new_para, para, last_run = self._handle_inline_image_item(item.image_url, current_item_level)
                continue

            # 表格
            if item.type == "table":
                need_new_para, para, last_run = self._handle_inline_table_item(item, current_item_level)
                continue

            # 行内代码
            if item.type == "inline_code":
                if need_new_para or para is None:
                    if need_new_para:
                        para = self._create_continuation_paragraph(list_type, level)
                        need_new_para = False
                    else:
                        para = self.document.add_paragraph()
                        if level > 0:
                            para.paragraph_format.left_indent = Inches(level * 0.5)
                    last_run = None
                    last_was_break = False
                self._add_inline_code_run(para, item)
                last_run = None  # 避免后续文本追加到代码 run 上
                continue

            # 代码块
            if item.type == "code":
                need_new_para, para, last_run = self._handle_inline_code_item(item, current_item_level)
                continue

    def _create_inline_paragraph(self, list_type: Optional[str], level: int):
        """
        创建内联段落；列表类型（有序/无序）均不预创建空段落，
        由 list_marker 处理器（_handle_list_marker_item）自行创建。

        Args:
            list_type: 列表类型（"ol" 或 "ul"）
            level: 嵌套层级

        Returns:
            Paragraph | None: 始终返回 None，由后续标记处理创建段落
        """
        return None

    def _create_continuation_paragraph(self, list_type: Optional[str], level: int):
        """创建列表项内部块级元素后的续写段落"""
        if list_type == "ul":
            para = self.document.add_paragraph(style="List Bullet")
        else:
            para = self.document.add_paragraph()
        if level > 0:
            para.paragraph_format.left_indent = Inches(level * 0.5)
        return para

    def _filter_unnecessary_linebreaks(self, items: list[InlineContent]) -> list[InlineContent]:
        """
        移除紧邻"会另起新段落的元素"之前的纯换行符。
        保留普通文本之间必要的换行。

        Args:
            items: 内联内容列表

        Returns:
            list[InlineContent]: 过滤后的内联内容列表
        """
        # 判断一个 item 是否会在 Word 中另起一个新段/块
        def _starts_new_block(item: InlineContent) -> bool:
            if item.list_marker:  # 嵌套列表项
                return True
            if item.type == "latex" and item.is_display:  # 块级公式
                return True
            if item.type in ("image", "table", "code"):  # 图片、表格、代码块
                return True
            return False

        filtered = []
        i = 0
        while i < len(items):
            item = items[i]
            # 当前是纯换行符？
            if item.type == "text" and item.content == "\n":
                # 找到下一个非换行的有效元素
                j = i + 1
                while j < len(items) and items[j].type == "text" and items[j].content == "\n":
                    j += 1
                # 如果下一个有效元素是需要另起段落的，丢弃所有连续换行
                if j < len(items) and _starts_new_block(items[j]):
                    i = j  # 跳过这些换行，直接处理后面的元素
                    continue
                else:
                    # 保留这些换行（后面是普通文本）
                    while i < j:
                        filtered.append(items[i])
                        i += 1
                    continue
            filtered.append(item)
            i += 1
        return filtered

    def _convert_latex_content(self, latex: str, is_display: bool) -> tuple[Optional[Element], str]:
        """
        转换 LaTeX 公式为 OMML 元素或 Unicode 回退文本

        Args:
            latex: LaTeX 公式文本（已去除界定符）
            is_display: 是否为展示公式

        Returns:
            tuple[Optional[Element], str]: (OMML 元素, Unicode 回退文本)，元素优先
        """
        latex = self._compensate_text_latex(latex)
        math_element = self._latex_to_omml(latex, is_display)
        if math_element is not None:
            return math_element, ""
        unicode_text = self.latex_converter.convert_inline(latex)
        return None, unicode_text

    def _add_latex(self, latex: str, is_display: bool = False) -> None:
        """
        添加 LaTeX 公式（独立成段）

        Args:
            latex: LaTeX 公式文本
            is_display: 是否为展示公式
        """
        if not latex.strip():
            return

        omml, unicode_text = self._convert_latex_content(latex, is_display)

        if omml is not None:
            try:
                p = self.document.add_paragraph()
                p.add_run(" ")
                p._element.append(deepcopy(omml))
                p.add_run(" ")
            except Exception as e:
                raise ExportError(f"添加公式失败: {e}") from e
        elif unicode_text:
            if is_display:
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(unicode_text)
                self._set_run_font(run)
            else:
                p = self.document.add_paragraph(unicode_text)
                self._set_run_font(p.runs[0])

    def _add_latex_to_paragraph(self, para, latex: str, is_display: bool = False) -> None:
        """
        添加内联公式到已有段落

        Args:
            para: 目标段落
            latex: LaTeX 公式文本
            is_display: 是否为展示公式
        """
        if not latex.strip():
            return

        omml, unicode_text = self._convert_latex_content(latex, is_display)

        if omml is not None:
            try:
                para.add_run(" ")
                para._element.append(deepcopy(omml))
                para.add_run(" ")
            except Exception as e:
                raise ExportError(f"添加公式失败: {e}") from e
        elif unicode_text:
            if is_display:
                p = self.document.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(unicode_text)
                self._set_run_font(run)
            else:
                run = para.add_run(unicode_text)
                self._set_run_font(run)

    def _latex_to_omml(self, latex: str, is_display: bool = False) -> Element | None:
        """
        使用 pandoc 将 LaTeX 转换为 OMML

        Args:
            latex: LaTeX 公式文本
            is_display: 是否为展示公式

        Returns:
            Element | None: OMML 元素，转换失败返回 None
        """
        deps_ok, _ = self.latex_converter.check_dependencies()  # pandoc 是否可用
        # pandoc 不可用时直接返回
        if not deps_ok:
            return None

        # 块级公式使用 $$...$$
        if is_display:
            tex_content = f'$$\n{latex}\n$$'
        # 行内公式使用 \(...\)
        else:
            tex_content = f'\\({latex}\\)'

        tmp_tex_path = None  # 临时 tex 文件路径
        tmp_docx_path = None  # 临时 docx 文件路径

        try:
            with tempfile.NamedTemporaryFile(
                mode='w', suffix='.tex', delete=False, encoding='utf-8'
            ) as tmp_tex:
                tmp_tex.write(tex_content)
                tmp_tex_path = tmp_tex.name

            with tempfile.NamedTemporaryFile(
                suffix='.docx', delete=False
            ) as tmp_docx:
                tmp_docx_path = tmp_docx.name

            cmd = ["pandoc", tmp_tex_path, "-o", tmp_docx_path]
            result = subprocess.run(cmd, capture_output=True, text=True, timeout=self._pandoc_timeout)

            # pandoc 转换成功：从生成的 docx 中提取公式
            if result.returncode == 0 and os.path.exists(tmp_docx_path):
                doc = Document(tmp_docx_path)
                for para in doc.paragraphs:
                    math_el = self._extract_math_from_paragraph(para._element)
                    if math_el is not None:
                        self._set_math_chinese_font(math_el, self.config.font_name)
                        return math_el
            if result.stderr:
                print(f"pandoc警告: {result.stderr[:200]}")
        except Exception as e:
            raise ExportError(f"Pandoc转换失败: {e}") from e
        finally:
            for path in [tmp_tex_path, tmp_docx_path]:
                if path:
                    try:
                        if os.path.exists(path):
                            os.unlink(path)
                    except OSError as e:
                        self._logger.debug(f"临时文件清理失败: {e}")

        return None

    def _extract_math_from_paragraph(self, p_element) -> Element | None:
        """
        从段落元素中提取数学公式

        Args:
            p_element: 段落 XML 元素

        Returns:
            Element | None: 数学公式元素
        """
        for child in p_element:  # 遍历段落子元素
            tag = child.tag  # 子元素标签
            if "oMath" in tag:
                return child
            if "oMathPara" in tag:
                return child
        return None

    def _set_math_chinese_font(self, math_element: Element, font_name: str) -> None:
        """
        设置数学公式中的中文字体

        Args:
            math_element: 数学公式 XML 元素
            font_name: 字体名称
        """
        W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"

        def qn_m(tag: str) -> str:  # 构建 Math 命名空间标签
            return "{%s}%s" % (M_NS, tag)

        def qn_w(tag: str) -> str:  # 构建 Word 命名空间标签
            return "{%s}%s" % (W_NS, tag)

        def is_chinese(text: str) -> bool:  # 判断是否包含中文字符
            return any('\u4e00' <= c <= '\u9fff' for c in text)

        for m_r in math_element.iter(qn_m("r")):
            m_t = m_r.find(qn_m("t"))
            if m_t is not None and m_t.text and is_chinese(m_t.text):
                w_rPr = m_r.find(qn_w("rPr"))
                if w_rPr is None:
                    from docx.oxml import parse_xml
                    w_rPr = parse_xml(
                        f'<w:rPr xmlns:w="{W_NS}">'
                        f'<w:rFonts w:hint="eastAsia" w:ascii="{font_name}" w:hAnsi="{font_name}" w:eastAsia="{font_name}" w:cs="{font_name}"/>'
                        f"<w:b w:val='0'/><w:i w:val='0'/>"
                        f"</w:rPr>"
                    )
                    m_r.append(w_rPr)

    def _compensate_text_latex(self, latex: str) -> str:
        """
        补偿 \\text{} 中的空格丢失问题，将所有前导空格替换为 ~

        Args:
            latex: LaTeX 公式文本

        Returns:
            str: 处理后的 LaTeX 文本
        """
        def _replace_spaces(match):
            """将连续空格替换为对应数量的 ~"""
            spaces = match.group(1)
            return '\\text{' + '~' * len(spaces)
        return re.sub(r'\\text\{( +)', _replace_spaces, latex)

    @staticmethod
    def _set_code_font(run) -> None:
        """
        为代码 run 设置等宽字体（西文 Courier New，东亚 SimSun）

        Args:
            run: docx Run 对象
        """
        rPr = run._element.get_or_add_rPr()
        rFonts = rPr.find(qn('w:rFonts'))
        if rFonts is None:
            rFonts = parse_xml(
                '<w:rFonts xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"/>'
            )
            rPr.append(rFonts)
        rFonts.set(qn('w:ascii'), 'Courier New')
        rFonts.set(qn('w:hAnsi'), 'Courier New')
        rFonts.set(qn('w:eastAsia'), 'SimSun')
        run.font.name = "Courier New"

    def _add_code_block(self, code: str, language: Optional[str] = None, level: int = 0) -> None:
        """
        添加代码块（使用 Courier New + 宋体 保证等宽对齐）
        """
        para = self.document.add_paragraph()
        if level > 0:
            para.paragraph_format.left_indent = Inches(level * 0.5)
        run = para.add_run(code)
        run.font.size = Pt(self.config.style_config.code_font_size if self.config.style_config else 10)
        self._set_code_font(run)

        # 添加灰色背景
        pPr = para._element.get_or_add_pPr()
        shd = parse_xml(r'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:val="pro" w:fill="F2F2F2"/>')
        pPr.append(shd)
        
    def _add_table(self, table_data) -> None:
        """
        添加表格

        Args:
            table_data: 表格数据对象

        Returns:
            None
        """
        if not table_data:
            return
        # 兜底：表头为空但存在数据行时，自动生成列号表头
        if not table_data.headers and table_data.rows:
            cols = len(table_data.rows[0]) if table_data.rows else 0
            table_data.headers = [f"列{i+1}" for i in range(cols)]
            table_data.header_bold = [False] * cols
        if not table_data.headers:
            return

        cols = len(table_data.headers)  # 列数
        rows_count = len(table_data.rows) + 1  # 行数（含表头）

        table = self.document.add_table(rows=rows_count, cols=cols)
        table.style = "Table Grid"

        # 添加表头
        header_bold = getattr(table_data, 'header_bold', []) or [False] * cols  # 表头加粗标记
        for i, header in enumerate(table_data.headers):
            cell = table.rows[0].cells[i]
            cell.text = header
            is_bold = header_bold[i] if i < len(header_bold) else True
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = is_bold
                    self._set_run_font(run)

        # 添加数据行
        cell_bold = getattr(table_data, 'cell_bold', []) or []  # 单元格加粗标记
        for row_idx, row_data in enumerate(table_data.rows):
            row_bold = cell_bold[row_idx] if row_idx < len(cell_bold) else []
            for col_idx, cell_data in enumerate(row_data):
                if col_idx < cols:
                    cell = table.rows[row_idx + 1].cells[col_idx]
                    cell.text = cell_data
                    is_bold = row_bold[col_idx] if col_idx < len(row_bold) else False
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = is_bold
                            self._set_run_font(run)

    def _add_blockquote(self, content: str) -> None:
        """
        添加引用块

        Args:
            content: 引用文本

        Returns:
            None
        """
        para = self.document.add_paragraph(content)
        para.style = "Quote"
        for run in para.runs:
            self._set_run_font(run)
            run.font.italic = True

    def _record_image_failure(self, url: str) -> None:
        """
        记录图片下载失败（去重）

        Args:
            url: 图片 URL
        """
        if url not in self._image_failure_urls_set:
            self._image_failure_urls_set.add(url)
            self._image_failure_count += 1
            self._image_failure_urls.append(url)

    @staticmethod
    def _is_image_content_type(content_type: str) -> bool:
        """
        判断 Content-Type 是否为图片类型

        Args:
            content_type: HTTP Content-Type 头

        Returns:
            bool: 是否为图片类型
        """
        if not content_type:
            return False
        ct_lower = content_type.lower()
        # 标准图片 MIME
        if ct_lower.startswith("image/"):
            return True
        # 常见二进制流类型（某些图床使用）
        if ct_lower in ("application/octet-stream", "binary/octet-stream"):
            return True
        return False

    @staticmethod
    def _has_image_extension(url: str) -> bool:
        """
        判断 URL 是否具有常见图片后缀

        Args:
            url: 图片 URL

        Returns:
            bool: 是否具有图片后缀
        """
        image_exts = (".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".svg", ".ico")
        return url.split("?")[0].lower().endswith(image_exts)

    @staticmethod
    def _is_image_magic(data: bytes) -> bool:
        """
        通过文件头魔数判断是否为常见图片格式

        Args:
            data: 文件二进制数据

        Returns:
            bool: 是否为已知图片格式
        """
        if len(data) < 8:
            return False
        # PNG
        if data[:8] == b'\x89PNG\r\n\x1a\n':
            return True
        # JPEG
        if data[:3] == b'\xff\xd8\xff':
            return True
        # GIF
        if data[:6] in (b'GIF89a', b'GIF87a'):
            return True
        # WebP
        if data[:4] == b'RIFF' and len(data) >= 12 and data[8:12] == b'WEBP':
            return True
        # BMP
        if data[:2] == b'BM':
            return True
        return False

    def _download_image(self, url: str) -> Optional[bytes]:
        """
        下载图片到内存

        Args:
            url: 图片 URL

        Returns:
            Optional[bytes]: 图片数据，下载失败返回 None
        """
        if not url or not url.startswith("http"):
            self._record_image_failure(url or "(空URL)")
            return None
        try:
            crawler_config = self._config.crawler
            user_agents = crawler_config.user_agents if crawler_config else None
            user_agent = user_agents[0] if user_agents else "Mozilla/5.0"
            resp = requests.get(url, timeout=self._image_timeout, headers={"User-Agent": user_agent})  # HTTP 响应
            resp.raise_for_status()

            content_type = resp.headers.get("content-type", "")  # 内容类型
            # Content-Type 非图片/二进制类型且 URL 无图片后缀时拒绝
            if content_type and not self._is_image_content_type(content_type):
                if not self._has_image_extension(url):
                    self._record_image_failure(url)
                    return None
            # 无 Content-Type 且 URL 无图片后缀时，进行魔数校验和 JSON/HTML 头检测
            if not content_type and not self._has_image_extension(url):
                if not self._is_image_magic(resp.content):
                    self._record_image_failure(url)
                    return None
                # 防御性检查：JSON/HTML 响应（以 { 或 < 开头）
                content = resp.content
                if content[:1] == b'{' or content[:1] == b'<':
                    self._record_image_failure(url)
                    return None
            return resp.content
        except requests.RequestException as e:
            self._record_image_failure(url)
            return None

    def _add_image(self, url: str, max_width: Optional[Inches] = None) -> None:
        """
        添加图片

        Args:
            url: 图片 URL
            max_width: 最大宽度

        Returns:
            None
        """
        image_data = self._download_image(url)
        if image_data is None:
            self._add_paragraph("[图片加载失败]")
            return

        temp_path = None
        try:
            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as f:
                f.write(image_data)
                temp_path = f.name
            width = max_width or Inches(self.config.style_config.image_width if self.config.style_config else 5.0)
            para = self.document.add_paragraph()
            para.add_run().add_picture(temp_path, width=width)
        except Exception as e:
            self._record_image_failure(url)
            self._add_paragraph(f"[图片: {url}]")
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.unlink(temp_path)
                except OSError:
                    pass
